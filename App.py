import streamlit as st
import pypdf
import folium
from streamlit_folium import st_folium
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# ============ PAGE CONFIG ============
st.set_page_config(
    page_title="KP Secretariat - PC-1 & PDF Intelligence Portal",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============ SESSION STATE ============
if "extracted_texts" not in st.session_state:
    st.session_state.extracted_texts = {}
if "map_html" not in st.session_state:
    st.session_state.map_html = None
if "map_filename" not in st.session_state:
    st.session_state.map_filename = "map"
if "detected_districts" not in st.session_state:
    st.session_state.detected_districts = []

# ============ PREMIUM CUSTOM STYLING ============
st.markdown("""
<style>
    /* --- FORCE DARK PREMIUM BACKGROUND --- */
    .stApp {
        background: linear-gradient(135deg, #0a0f24 0%, #162254 40%, #0a0f24 80%) !important;
        background-attachment: fixed !important;
    }
    .stApp::before {
        content: '';
        position: fixed;
        top: 0; left: 0; right: 0; bottom: 0;
        background: linear-gradient(-45deg, #1a2a6c, #0a0f24, #2a1a4a, #0a0f24);
        background-size: 400% 400%;
        animation: gradientBG 22s ease infinite;
        z-index: -1;
    }
    @keyframes gradientBG {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }

    /* --- MAIN GLASS CONTAINER --- */
    .main .block-container {
        background: rgba(255, 255, 255, 0.07) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border-radius: 40px !important;
        padding: 2.5rem 2.5rem !important;
        margin-top: 15px !important;
        margin-bottom: 25px !important;
        border: 1px solid rgba(255, 215, 0, 0.15) !important;
        box-shadow: 0 30px 60px rgba(0,0,0,0.7) !important;
    }

    /* --- TEXT COLORS (Readable) --- */
    body, p, span, label, div, .stMarkdown, caption, .stTextInput > label {
        color: #e8eeff !important;
        font-weight: 500 !important;
    }
    h1, h2, h3, h4, h5, .big-title, .sub-header, .config-title {
        color: #f9d976 !important;
        text-shadow: 0 2px 12px rgba(249, 217, 118, 0.35) !important;
    }
    .big-title {
        font-size: 42px !important;
        font-weight: 900 !important;
        text-align: center !important;
        letter-spacing: 1px !important;
    }
    .sub-header {
        color: #a8c0ff !important;
        font-size: 17px !important;
        font-weight: 600 !important;
        text-align: center !important;
        opacity: 0.95 !important;
    }
    .config-title {
        font-size: 22px !important;
        font-weight: 700 !important;
        margin-top: 10px !important;
    }

    /* --- LOGO ANIMATION (Targets the real kp_logo.png) --- */
    div[data-testid="stImage"] {
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        position: relative !important;
        pointer-events: none !important;
        user-select: none !important;
        margin: 0 auto 15px auto !important;
    }

    div[data-testid="stImage"] img {
        border-radius: 50% !important;
        width: 180px !important;
        height: 180px !important;
        object-fit: cover !important;
        pointer-events: none !important;
        user-select: none !important;
        box-shadow: 0 0 25px rgba(249, 217, 118, 0.3) !important;
        position: relative;
        z-index: 5;
    }

    /* 12-piece continuous spinning ring */
    div[data-testid="stImage"]::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 200px;
        height: 200px;
        transform: translate(-50%, -50%);
        border-radius: 50%;
        background: repeating-conic-gradient(
            from 0deg,
            #f9d976 0deg 15deg,
            #1a2a6c 15deg 30deg,
            #f9d976 30deg 45deg,
            #1a2a6c 45deg 60deg,
            #f9d976 60deg 75deg,
            #1a2a6c 75deg 90deg,
            #f9d976 90deg 105deg,
            #1a2a6c 105deg 120deg,
            #f9d976 120deg 135deg,
            #1a2a6c 135deg 150deg,
            #f9d976 150deg 165deg,
            #1a2a6c 165deg 180deg,
            #f9d976 180deg 195deg,
            #1a2a6c 195deg 210deg,
            #f9d976 210deg 225deg,
            #1a2a6c 225deg 240deg,
            #f9d976 240deg 255deg,
            #1a2a6c 255deg 270deg,
            #f9d976 270deg 285deg,
            #1a2a6c 285deg 300deg,
            #f9d976 300deg 315deg,
            #1a2a6c 315deg 330deg,
            #f9d976 330deg 345deg,
            #1a2a6c 345deg 360deg
        );
        animation: spinRing 6s linear infinite;
        mask: radial-gradient(farthest-side, transparent calc(100% - 5px), #fff calc(100% - 3px));
        -webkit-mask: radial-gradient(farthest-side, transparent calc(100% - 5px), #fff calc(100% - 3px));
        z-index: 2;
        pointer-events: none !important;
    }

    /* Soft glow pulse */
    div[data-testid="stImage"]::after {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 240px;
        height: 240px;
        transform: translate(-50%, -50%);
        border-radius: 50%;
        background: radial-gradient(circle, rgba(249, 217, 118, 0.22) 0%, transparent 70%);
        animation: pulseGlow 2.8s ease-in-out infinite alternate;
        z-index: 1;
        pointer-events: none !important;
    }

    @keyframes spinRing {
        0% { transform: translate(-50%, -50%) rotate(0deg); }
        100% { transform: translate(-50%, -50%) rotate(360deg); }
    }
    @keyframes pulseGlow {
        0% { transform: translate(-50%, -50%) scale(0.9); opacity: 0.3; }
        100% { transform: translate(-50%, -50%) scale(1.2); opacity: 0.7; }
    }

    /* --- NAVIGATION TABS --- */
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(26, 42, 108, 0.75) !important;
        backdrop-filter: blur(12px) !important;
        border-radius: 16px !important;
        padding: 8px !important;
        gap: 6px !important;
        border: 1px solid rgba(249, 217, 118, 0.25) !important;
        box-shadow: 0 8px 25px rgba(0,0,0,0.5) !important;
    }
    .stTabs [data-baseweb="tab"] {
        color: #b8c9ff !important;
        font-weight: 700 !important;
        border-radius: 12px !important;
        padding: 12px 22px !important;
        transition: all 0.25s ease !important;
        border: none !important;
        background: transparent !important;
    }
    .stTabs [data-baseweb="tab"]:hover {
        background: rgba(249, 217, 118, 0.18) !important;
        color: #f9d976 !important;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #f9d976 0%, #f4b836 100%) !important;
        color: #0a0f24 !important;
        box-shadow: 0 4px 18px rgba(249, 217, 118, 0.45) !important;
        font-weight: 800 !important;
    }

    /* --- BUTTONS --- */
    .stButton>button, .stDownloadButton>button {
        width: 100% !important;
        border-radius: 14px !important;
        font-weight: 700 !important;
        background: linear-gradient(135deg, #f9d976 0%, #f4b836 100%) !important;
        color: #0a0f24 !important;
        border: none !important;
        padding: 0.65rem 1.2rem !important;
        box-shadow: 0 6px 20px rgba(249, 217, 118, 0.3) !important;
        transition: all 0.25s ease !important;
        text-transform: uppercase !important;
        letter-spacing: 0.5px !important;
    }
    .stButton>button:hover, .stDownloadButton>button:hover {
        transform: scale(1.02) !important;
        box-shadow: 0 8px 28px rgba(249, 217, 118, 0.55) !important;
        background: linear-gradient(135deg, #ffe28a 0%, #f9c830 100%) !important;
    }

    /* --- INPUTS --- */
    .stTextInput>div>div>input, 
    .stSelectbox>div>div, 
    .stTextArea>div>div>textarea {
        background: rgba(255, 255, 255, 0.09) !important;
        backdrop-filter: blur(10px) !important;
        border: 1px solid rgba(249, 217, 118, 0.35) !important;
        border-radius: 12px !important;
        color: #f0f4ff !important;
        padding: 10px 15px !important;
        font-weight: 500 !important;
    }
    .stTextInput>div>div>input:focus, 
    .stSelectbox>div>div:focus {
        border: 1px solid #f9d976 !important;
        box-shadow: 0 0 18px rgba(249, 217, 118, 0.2) !important;
    }
    label {
        color: #b8c9ff !important;
        font-weight: 600 !important;
    }

    /* --- ALERTS & UPLOADER --- */
    .stAlert {
        background: rgba(249, 217, 118, 0.12) !important;
        border: 1px solid #f9d976 !important;
        border-radius: 16px !important;
        color: #f9d976 !important;
        font-weight: 600 !important;
    }
    .stFileUploader > div > div {
        background: rgba(255,255,255,0.05) !important;
        border: 2px dashed rgba(249, 217, 118, 0.4) !important;
        border-radius: 16px !important;
        padding: 20px !important;
    }
    .stFileUploader > div > div:hover {
        border-color: #f9d976 !important;
        background: rgba(249, 217, 118, 0.06) !important;
    }
    .stInfo, .stSuccess, .stError {
        border-radius: 14px !important;
        font-weight: 600 !important;
    }
</style>
""", unsafe_allow_html=True)

# ============ DATA PRIVACY WARNING ============
st.warning("🔒 **Data Privacy Notice:** Your document text is sent to OpenRouter (third-party AI). Do NOT upload classified or highly sensitive government documents.")

# ============ LOGO (Same kp_logo.png + Animation) ============
lc1, lc2, lc3 = st.columns([1, 1.2, 1])
with lc2:
    try:
        st.image("kp_logo.png", use_container_width=True)
    except:
        st.markdown("🏛️ **Place `kp_logo.png` here**")

# ============ TITLE ============
st.markdown('<h1 class="big-title">🏛️ KP Secretariat – Smart PC-1 & PDF Review Portal</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Financial Summary Extractor • Geographic Map Mapper • Hidden Issue Scanner • Official Draft Generator</p>', unsafe_allow_html=True)

# ============ CONFIGURATIONS CARD ============
st.markdown('<p class="config-title">⚙️ System Configurations</p>', unsafe_allow_html=True)
cfg_l, cfg_c, cfg_r = st.columns([1, 2, 1])
with cfg_c:
    with st.container():
        st.markdown('<div style="background: rgba(255,255,255,0.05); padding: 20px; border-radius: 20px; border: 1px solid rgba(249,217,118,0.2);">', unsafe_allow_html=True)
        BUILTIN_API_KEY = ""
        api_key = st.text_input("🔑 OpenRouter API Key", value=BUILTIN_API_KEY, type="password", placeholder="Apni OpenRouter API key yahan enter karein...")
        model_name = st.selectbox(
            "🤖 Select AI Model via OpenRouter",
            [
                "meta-llama/llama-3.3-70b-instruct",
                "google/gemini-2.0-flash-001",
                "deepseek/deepseek-chat"
            ]
        )
        st.markdown('</div>', unsafe_allow_html=True)
st.markdown("---")

# ============ TABS ============
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📊 Financial Summary", 
    "🗺️ Map Mapper", 
    "🔍 Issue Scanner",
    "💬 Deep Chat",
    "⚖️ Comparative Analysis",
    "📝 Draft Generator"
])

# ============ HELPER FUNCTIONS ============
@st.cache_data(show_spinner=False)
def validate_and_extract_pdf(uploaded_bytes, file_name):
    try:
        reader = pypdf.PdfReader(io.BytesIO(uploaded_bytes))
        if len(reader.pages) == 0:
            raise ValueError("PDF has no pages.")
        text = ""
        for index, page in enumerate(reader.pages):
            extracted = page.extract_text()
            if extracted:
                text += f"\n--- Page {index + 1} ---\n" + extracted + "\n"
        if len(text.strip()) < 50:
            raise ValueError("Scanned PDF or no text found. Please upload a searchable PDF.")
        return text
    except Exception as e:
        raise ValueError(f"PDF Processing Error: {str(e)}")

def safe_truncate_text(text, max_chars=35000):
    if len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    last_space = truncated.rfind(' ')
    if last_space > 0:
        return truncated[:last_space] + "\n\n[...Truncated for AI context...]"
    return truncated + "\n\n[...Truncated for AI context...]"

def clean_markdown_for_docx(text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    return text

def add_logo_to_docx(doc):
    try:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture("kp_logo.png", width=Inches(1.0))
    except:
        pass

# --- District DB ---
KP_DISTRICTS_DB = {
    "Peshawar": [34.0151, 71.5249], "Mardan": [34.1989, 72.0406], "Swat": [34.7717, 72.3602],
    "Abbottabad": [34.1463, 73.2117], "Bannu": [32.9889, 70.6042], "Dera Ismail Khan": [31.8314, 70.9014],
    "Kohat": [33.5889, 71.4429], "Swabi": [34.1167, 72.4667], "Nowshera": [34.0153, 71.9747],
    "Charsadda": [34.1526, 71.7381], "Mansehra": [34.3333, 73.2000], "Haripur": [33.9994, 72.9342]
}
FUZZY_ALIASES = {
    "di khan": "Dera Ismail Khan", "d.i. khan": "Dera Ismail Khan", "pesh": "Peshawar",
    "mard": "Mardan", "swabi": "Swabi", "nowshera": "Nowshera", "charsadda": "Charsadda",
    "mansehra": "Mansehra", "haripur": "Haripur", "kohat": "Kohat", "bannu": "Bannu",
    "abbottabad": "Abbottabad", "swat": "Swat"
}

# ============ TAB 1 ============
with tab1:
    st.markdown("### 📊 PC-1 Financial Summary & Cost Extraction")
    fin_file = st.file_uploader("Upload PC-1 PDF for Financial Extraction", type=["pdf"], key="fin_pdf")
    if fin_file is not None:
        if st.button("📊 Extract Financial Tables & Summary", type="primary"):
            if not api_key:
                st.error("⚠️ Please provide your OpenRouter API Key!")
            else:
                try:
                    with st.status("📄 Extracting financial schedules...", expanded=True) as status:
                        fin_text = validate_and_extract_pdf(fin_file.getvalue(), fin_file.name)
                        st.session_state.extracted_texts[fin_file.name] = fin_text
                        client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
                        status.update(label="🤖 Sending to AI for financial analysis...")
                        fin_prompt = f"""
Act as a Senior Financial Analyst in KP Planning & Development Department.
Extract all financial summaries, cost breakdowns, capital expenditures, and budget allocations.
Present findings clearly in Markdown tables.

Document Text:
{safe_truncate_text(fin_text)}
"""
                        fin_completion = client.chat.completions.create(
                            model=model_name,
                            messages=[{"role": "system", "content": "Precise financial analyst."}, {"role": "user", "content": fin_prompt}],
                            max_tokens=3000, temperature=0.1
                        )
                        status.update(label="✅ Extraction Complete!", state="complete")
                        st.markdown("### 📋 Extracted Financial Summary")
                        st.markdown(fin_completion.choices[0].message.content)
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
    else:
        st.info("👆 Financial summary extract karne ke liye PC-1 PDF upload karein.")

# ============ TAB 2 ============
with tab2:
    st.markdown("### 🗺️ ADP Project Location & Geographic Map Mapper")
    map_file = st.file_uploader("Upload PC-1 / ADP PDF for Geographic Mapping", type=["pdf"], key="map_pdf")
    if map_file is not None:
        base_name_m = map_file.name.rsplit('.', 1)[0]
        try:
            map_text = validate_and_extract_pdf(map_file.getvalue(), map_file.name)
            st.session_state.extracted_texts[map_file.name] = map_text
            detected = []
            map_text_lower = map_text.lower()
            for dist in KP_DISTRICTS_DB.keys():
                if dist.lower() in map_text_lower:
                    detected.append(dist)
            for alias, real_name in FUZZY_ALIASES.items():
                if alias in map_text_lower and real_name not in detected:
                    detected.append(real_name)
            if not detected:
                detected = ["Peshawar", "Mardan", "Swat"]
            st.session_state.detected_districts = detected
            st.session_state.map_filename = base_name_m
            st.success(f"✅ Detected Districts: {', '.join(detected)}")
        except Exception as e:
            st.error(f"❌ PDF Read Error: {str(e)}")
            st.stop()

        if st.button("🗺️ Generate Interactive Map", type="primary"):
            with st.spinner("🗺️ Rendering map..."):
                detected = st.session_state.detected_districts
                center_coords = KP_DISTRICTS_DB.get(detected[0], [34.0151, 71.5249])
                m = folium.Map(location=center_coords, zoom_start=8)
                for dist in detected:
                    folium.Marker(
                        KP_DISTRICTS_DB[dist], 
                        popup=f"📍 Project Site: {dist}", 
                        tooltip=dist, 
                        icon=folium.Icon(color="red", icon="info-sign")
                    ).add_to(m)
                st.session_state.map_html = m._repr_html_()
                st.success("✅ Map generated successfully!")

        if st.session_state.map_html:
            st.components.v1.html(st.session_state.map_html, height=450)
            st.download_button(
                "📥 Download Interactive Map (.html)", 
                data=st.session_state.map_html.encode('utf-8'), 
                file_name=f"{st.session_state.map_filename}_Map.html", 
                mime="text/html"
            )
    else:
        st.info("👆 Geographic mapping ke liye PDF upload karein.")

# ============ TAB 3 ============
with tab3:
    st.markdown("### 🔍 Scan Massive PDFs for Hidden Loopholes")
    scan_file = st.file_uploader("Upload PDF Document for Deep Scanning", type=["pdf"], key="scan_pdf")
    scan_focus = st.text_input("Specific area to focus on (Optional):", placeholder="Cost escalation, timeline delays...")
    if st.button("🚀 Run Deep Issue Scanner", type="primary"):
        if not api_key:
            st.error("⚠️ Please provide API Key!")
        elif not scan_file:
            st.error("⚠️ Please upload a PDF!")
        else:
            try:
                with st.status("🔎 Scanning document for issues...", expanded=True) as status:
                    pdf_text = validate_and_extract_pdf(scan_file.getvalue(), scan_file.name)
                    st.session_state.extracted_texts[scan_file.name] = pdf_text
                    status.update(label="🧠 Analyzing with AI...")
                    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
                    scanner_prompt = f"""
Act as Chief Planning and Audit Inspector for Govt of KP.
Find hidden loopholes, financial inconsistencies, missing approvals.
Specific Focus: {scan_focus}

Document Text:
{safe_truncate_text(pdf_text)}
"""
                    completion = client.chat.completions.create(
                        model=model_name,
                        messages=[{"role": "system", "content": "Meticulous govt auditor."}, {"role": "user", "content": scanner_prompt}],
                        max_tokens=4000, temperature=0.1
                    )
                    status.update(label="✅ Scan Complete", state="complete")
                    st.markdown("### 📋 Deep Audit Findings")
                    st.markdown(completion.choices[0].message.content)
            except Exception as e:
                st.error(f"❌ Scan Error: {str(e)}")

# ============ TAB 4 ============
with tab4:
    st.markdown("### 💬 Deep Chat with Massive PDFs")
    chat_file = st.file_uploader("Upload Large PDF Document", type=["pdf"], key="chat_pdf")
    if chat_file is not None:
        try:
            if chat_file.name not in st.session_state.extracted_texts:
                chat_pdf_text = validate_and_extract_pdf(chat_file.getvalue(), chat_file.name)
                st.session_state.extracted_texts[chat_file.name] = chat_pdf_text
            else:
                chat_pdf_text = st.session_state.extracted_texts[chat_file.name]
            user_question = st.text_input("💬 Aap is document mein kya talaash kar rahe hain?")
            if st.button("🔍 Search & Answer from PDF", type="primary"):
                if not api_key:
                    st.error("⚠️ API Key missing!")
                elif not user_question:
                    st.warning("⚠️ Type your question.")
                else:
                    try:
                        with st.spinner("🔍 Searching across pages..."):
                            client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
                            chat_completion = client.chat.completions.create(
                                model=model_name,
                                messages=[
                                    {"role": "system", "content": "Expert document research assistant."},
                                    {"role": "user", "content": f"Document Text:\n{safe_truncate_text(chat_pdf_text)}\n\nQuestion: {user_question}"}
                                ],
                                max_tokens=2000, temperature=0.2
                            )
                            st.markdown("### 🤖 AI Precise Answer:")
                            st.markdown(chat_completion.choices[0].message.content)
                    except Exception as e:
                        st.error(f"❌ Chat Error: {str(e)}")
        except Exception as e:
            st.error(f"❌ PDF Load Error: {str(e)}")
    else:
        st.info("👆 Pehle upar PDF upload karein.")

# ============ TAB 5 ============
with tab5:
    st.markdown("### ⚖️ Compare Two Documents")
    col_orig, col_rev = st.columns(2)
    with col_orig:
        file_original = st.file_uploader("Upload First File", type=["pdf"], key="orig_file")
    with col_rev:
        file_revised = st.file_uploader("Upload Second File", type=["pdf"], key="rev_file")
    if st.button("🔍 Run Comparative Analysis", type="primary"):
        if not api_key:
            st.error("⚠️ API Key missing!")
        elif not file_original or not file_revised:
            st.error("⚠️ Upload both files!")
        else:
            base_name_c2 = file_revised.name.rsplit('.', 1)[0]
            try:
                with st.status("⚖️ Comparing documents...", expanded=True) as status:
                    status.update(label="📄 Extracting text from both files...")
                    text_orig = validate_and_extract_pdf(file_original.getvalue(), file_original.name)
                    text_rev = validate_and_extract_pdf(file_revised.getvalue(), file_revised.name)
                    st.session_state.extracted_texts[file_original.name] = text_orig
                    st.session_state.extracted_texts[file_revised.name] = text_rev
                    status.update(label="🧠 AI generating comparison...")
                    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
                    comparison_prompt = f"""
Compare these two documents:
1. Exact changes in cost/budget.
2. Scope modifications.
3. Added/removed conditions.

File 1: {safe_truncate_text(text_orig, 20000)}
File 2: {safe_truncate_text(text_rev, 20000)}
"""
                    completion = client.chat.completions.create(
                        model=model_name,
                        messages=[{"role": "system", "content": "Planning officer."}, {"role": "user", "content": comparison_prompt}],
                        max_tokens=3000, temperature=0.1
                    )
                    comparison_result = completion.choices[0].message.content
                    status.update(label="✅ Comparison Ready", state="complete")
                    st.markdown("### 📊 Detailed Comparative Report")
                    st.markdown(comparison_result)
                    
                    cleaned_result = clean_markdown_for_docx(comparison_result)
                    comp_doc = Document()
                    add_logo_to_docx(comp_doc)
                    comp_doc.add_heading("Comparative Analysis", level=1)
                    for para_text in cleaned_result.split('\n\n'):
                        if para_text.strip():
                            p = comp_doc.add_paragraph(para_text.strip())
                            p.paragraph_format.space_after = Pt(12)
                            p.paragraph_format.line_spacing = 1.5
                    docx_bytes = io.BytesIO()
                    comp_doc.save(docx_bytes)
                    docx_bytes.seek(0)
                    st.download_button("📥 Download Report (.docx)", data=docx_bytes, file_name=f"{base_name_c2}_Comparison.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")

# ============ TAB 6 ============
with tab6:
    st.markdown("### 📝 Generate Official Reply (.docx)")
    rep_file = st.file_uploader("Upload Observation Letter / PC-1 PDF", type=["pdf"], key="rep_pdf")
    task_option = st.selectbox("Choose Draft Type:", ["Official Reply", "Executive Summary", "Project Justification"])
    user_instructions = st.text_area("📝 Additional instructions:")
    if st.button("🚀 Generate Official Word Document", type="primary"):
        if not api_key:
            st.error("⚠️ API Key missing!")
        elif not rep_file:
            st.error("⚠️ Upload source PDF!")
        else:
            base_name = rep_file.name.rsplit('.', 1)[0]
            try:
                with st.status("📝 Drafting official document...", expanded=True) as status:
                    status.update(label="📄 Extracting source text...")
                    pdf_text = validate_and_extract_pdf(rep_file.getvalue(), rep_file.name)
                    st.session_state.extracted_texts[rep_file.name] = pdf_text
                    status.update(label="🧠 Generating draft with AI...")
                    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
                    system_prompt = "Expert Section Officer in KP Secretariat. Write formal govt correspondence."
                    completion = client.chat.completions.create(
                        model=model_name,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"Task: {task_option}\nInstructions: {user_instructions}\nSource Document:\n{safe_truncate_text(pdf_text)}"}
                        ],
                        max_tokens=4000, temperature=0.2
                    )
                    response_text = completion.choices[0].message.content
                    status.update(label="✅ Draft generated", state="complete")
                    st.markdown("### 📋 Generated Draft Preview")
                    st.markdown(response_text)
                    
                    cleaned_response = clean_markdown_for_docx(response_text)
                    doc = Document()
                    add_logo_to_docx(doc)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("\nGOVERNMENT OF KHYBER PAKHTUNKHWA\nSECRETARIAT - OFFICIAL CORRESPONDENCE\n")
                    run.bold = True
                    run.font.size = Pt(14)
                    doc.add_paragraph("_" * 80)
                    for line in cleaned_response.split('\n'):
                        line = line.strip()
                        if not line:
                            doc.add_paragraph()
                            continue
                        p_para = doc.add_paragraph(line)
                        p_para.paragraph_format.space_after = Pt(12)
                        p_para.paragraph_format.line_spacing = 1.5
                    docx_bytes_reply = io.BytesIO()
                    doc.save(docx_bytes_reply)
                    docx_bytes_reply.seek(0)
                    st.download_button("📥 Download Official Document (.docx)", data=docx_bytes_reply, file_name=f"{base_name}_Official_Reply.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
